import streamlit as st
import pandas as pd
from docx import Document
import io
import os

# ==========================================
# 設定エリア：ファイル名を実際の環境に合わせること
# ==========================================
EXCEL_FILE = "TOP-046付表 化学薬品台帳_TEST.xlsx" 
SHEET_NAME = "薬品・油剤台帳_新"
WORD_TEMPLATE = "TOP-046付表_MSDS作成用_TEST.docx" # 必ず.docx形式にすること
# ==========================================

@st.cache_data
def load_data():
    if not os.path.exists(EXCEL_FILE):
        return pd.DataFrame()
    
    # Excelの3行目（index=2）が見出し行であるため、header=2を指定
    df = pd.read_excel(EXCEL_FILE, sheet_name=SHEET_NAME, header=2, dtype=str).fillna("")
    
    # 複数行の見出しによる列名の改行ノイズ（「海洋汚染\n防止法」など）を消去
    df.columns = df.columns.str.replace('\n', '', regex=False)
    return df

def create_report(template_path, data_row):
    doc = Document(template_path)
    
    # --- 適用法規の動的結合処理（空白を飛ばすロジック） ---
    law_columns = ['消防法', '毒、劇物取締法', '労働安全衛生法', 'PRTR法', '水質汚濁防止法', 
                   '大気汚染', '水道法', '下水道法', '産廃・清掃法', '海洋汚染防止法', '航空法', '有機溶剤']
    
    applicable_laws = []
    for col in law_columns:
        if col in data_row and str(data_row[col]).strip() != "":
            # 値が存在する場合のみリストに追加
            applicable_laws.append(f"■ {col} : {data_row[col]}")
    
    # リストを改行で結合（何もなければ「該当法規なし」とする）
    law_text = "\n".join(applicable_laws) if applicable_laws else "該当法規なし"

    # --- Wordのタグと、流し込むテキストの紐付け辞書 ---
    replace_dict = {
        "品名・型式・規格": data_row.get("品名・型式・規格", ""),
        "適用法規": law_text,
        "有害性": data_row.get("有害性", ""),
        "吸入": data_row.get("吸入", ""),
        "皮膚": data_row.get("皮膚", ""),
        "目": data_row.get("目", ""),
        "飲込": data_row.get("飲込", ""),
        "消火剤": data_row.get("消火剤", ""),
        "特定の消火方法": data_row.get("特定の消火方法", ""),
        "消火者の保護等": data_row.get("消火者の保護等", ""),
        "取扱注意": data_row.get("取扱注意", ""),
        "保管注意": data_row.get("保管注意", "")
    }

    # --- 置換実行（段落テキスト） ---
    for p in doc.paragraphs:
        for key, value in replace_dict.items():
            tag = f"{{{{{key}}}}}"
            if tag in p.text:
                p.text = p.text.replace(tag, str(value))
    
    # --- 置換実行（表の中のテキスト） ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in replace_dict.items():
                        tag = f"{{{{{key}}}}}"
                        if tag in p.text:
                            p.text = p.text.replace(tag, str(value))

    # 完成したWordをメモリ上に保持
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# UI構築（Streamlit）
# ==========================================
st.set_page_config(page_title="SDS管理システム", layout="wide")
st.title("🧪 SDS管理・社内掲示MSDS自動生成")

df = load_data()

if df.empty:
    st.error(f"データベース ({EXCEL_FILE}) が見つからないか、読み込めません。")
    st.stop()

# 空白行（品名がない行）を視覚から除外
if "品名・型式・規格" in df.columns:
    df = df[df["品名・型式・規格"] != ""]

# 1. 検索機能
st.sidebar.header("🔍 検索")
search_term = st.sidebar.text_input("品名で検索")

if search_term:
    mask = df["品名・型式・規格"].astype(str).str.contains(search_term, case=False)
    filtered_df = df[mask]
else:
    filtered_df = df

st.dataframe(filtered_df, use_container_width=True)

st.divider()

# 2. 帳票出力機能
st.header("📄 掲示用MSDSの出力")

if not filtered_df.empty:
    selected_item = st.selectbox("出力する物質を選択", filtered_df["品名・型式・規格"].tolist())

    if selected_item:
        # 選択された行のデータを取得
        target_data = filtered_df[filtered_df["品名・型式・規格"] == selected_item].iloc[0]
        
        if os.path.exists(WORD_TEMPLATE):
            if st.button("🖨️ A4掲示用Wordファイルを生成", type="primary"):
                report_bytes = create_report(WORD_TEMPLATE, target_data)
                
                # ファイル名に使えない記号を置換
                safe_name = str(selected_item).replace("/", "_").replace("\\", "_")
                
                st.download_button(
                    label="📥 ダウンロード",
                    data=report_bytes,
                    file_name=f"MSDS_{safe_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error(f"⚠️ テンプレートファイル ({WORD_TEMPLATE}) が同じフォルダに見つかりません。")
else:
    st.info("該当する物質がありません。")